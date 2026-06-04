import os
import imaplib
import email
import json
import re
import smtplib
import requests
import io as _io_module
from email import encoders as _email_encoders
import redis as _redis_lib
from email.header import decode_header
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from datetime import datetime, timedelta
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, JobQueue
import anthropic

TELEGRAM_TOKEN = os.environ['TELEGRAM_TOKEN']
CHAT_ID = int(os.environ['CHAT_ID'])

# Contas monitoradas
ACCOUNTS = [
    {
        'user': os.environ['GMAIL_USER'],
        'password': os.environ['GMAIL_APP_PASSWORD'],
        'label': '💼 Trabalho',
    },
    {
        'user': os.environ.get('GMAIL_USER_2', ''),
        'password': os.environ.get('GMAIL_APP_PASSWORD_2', ''),
        'label': '👤 Pessoal',
    },
]
ACCOUNTS = [a for a in ACCOUNTS if a['user'] and a['password']]

# Compat: manter GMAIL_USER/GMAIL_APP_PASSWORD apontando para a primeira conta
GMAIL_USER = ACCOUNTS[0]['user']
GMAIL_APP_PASSWORD = ACCOUNTS[0]['password']

# ─── ALERTAS IMEDIATOS ────────────────────────────────────────────────────────
# Palavras que disparam 🚨 ALERTA na hora, independente do horário

ALERT_KEYWORDS = [
    # Pessoas chave
    'fellipe guerra', 'vanessa', 'ipog',
    # Obrigações fiscais
    'sped', 'sped fiscal', 'sped contábil', 'sped contabil', 'efd',
    'ecf', 'ecd', 'reinf', 'esocial', 'e-social', 'dctf', 'gia', 'gim',
    'nfe', 'nf-e', 'nota fiscal', 'xml', 'danfe',
    # Clientes (nomes distintos da planilha)
    'flex wind', 'guerra advogados', 'guerra treinamentos',
    'growup', 'italtecnology', 'iqnus', 'kgary', 'gmad',
    'milk heroes', 'mlvl', 'mourao estrutura', 'novo frio',
    'fit & fresh', 'fg coworking', 'fg consultoria',
    'hbr logistica', 'heverline', 'gutemberg',
    'farmacia filadelfia', 'farmacia eri', 'm3 farmacia',
    'lead imobiliaria', 'loteamento portal',
    'jessica furtado', 'luana lima', 'ivone cordeiro',
    'franklin cunha', 'marluce pinheiro',
    'c4 comercio', 'c4 consultoria', 'az servicos',
    'instituto ect', 'instituto de educacao contabil',
    'deusamim', 'vg braga', 'braga reis',
    '4 estylos', 'erico brasil',
    # Remetentes específicos monitorados
    'cfla', 'cfla-intl', 'serviceclient@cfla', 't1134',
    'raquel brito', 'quellbrito@hotmail', 'rafaelabbrito',
    # Ações urgentes
    'intimação', 'intimacao', 'auto de infração', 'auto de infracao',
    'embargo', 'autuação', 'autuacao', 'notificação fiscal', 'notificacao fiscal',
    'multa de ofício', 'multa de oficio', 'parcelamento', 'execução fiscal',
    'penhora', 'bloqueio judicial', 'certidão negativa', 'certidao negativa',
    'regularize', 'pgfn', 'dívida ativa', 'divida ativa',
]

# ─── FILTROS ──────────────────────────────────────────────────────────────────

SPAM_SENDERS = [
    'netshoes', 'nordresearch', 'avast', 'cvc.com', 'skyscanner',
    'grancursosonline', 'serasa', 'smiles.com', 'facebookmail',
    'insiderstore', 'wispr.ai', 'academia-mail', 'hotmilhas',
    'retornar.com', 'accor', 'reminders@', 'news.all@mail',
    'grancursos', 'empiricus', 'infomoney',
]
SPAM_SUBJECTS = [
    '% off', 'cupom', 'oferta imperd', 'passagem', 'megapromo',
    'economize', 'desconto extra', 'prorrogamos', 'feirao', 'feirão',
    'voucher', 'milhas', 'resort', 'black friday',
    'ultima chance', 'última chance', 'imperdivel', 'imperdível',
]
URGENT_SENDERS = [
    'atendimento@compliance-ce.com.br', 'pgfn.gov.br', 'qive.com.br',
    'sigmavaf.com.br', 'accounts.google.com', 'adveronix.com',
    'regularize', 'receita.fazenda', 'flex-wind.com', 'tailscale',
    'sefaz', 'prefeitura', 'tce.', 'tribunal',
]
URGENT_SUBJECTS = [
    'gestta', 'vencer', 'certificado', 'regularize', 'seguranca',
    'segurança', 'prefeitura', 'pgfn', 'multa', 'vencimento',
    'prazo', 'pendencia', 'pendência', 'provisao', 'provisão',
    'reversao', 'reversão', 'trial ends', 'notificacao', 'notificação',
    'intimacao', 'intimação', 'auto de infração',
]
SALES_SENDERS = ['eduzz.com', 'hotmart.com']
PROFESSIONAL_SENDERS = [
    'conjur.com.br', 'legisweb.com.br', 'reformatributaria.com.br', 'fbc.org.br',
]
PERSONAL_SENDERS = ['colmaster.com.br', 'isabelbtz@gmail']
FINANCIAL_SENDERS = ['xpi.com.br', 'xpi.com', 'nubank', 'itau', 'bradesco', 'bb.com.br']

ICONS = {
    'URGENTE': '🔴',
    'VENDAS': '💰',
    'PROFISSIONAL': '👨‍💼',
    'PESSOAL': '👦',
    'FINANCEIRO': '💳',
    'OUTROS': '📌',
}

# ─── CLAUDE AGENT — TOOL USE ──────────────────────────────────────────────────

AGENT_SYSTEM = """Você é o assistente pessoal de Marcos Lima, contador tributário da Compliance-CE (marcoslima@compliance-ce.com.br).
Você opera dentro do Telegram e tem acesso às ferramentas abaixo para gerenciar emails, gerar pareceres tributários e criar cards no Trello.

Regras:
- Responda sempre em português, de forma direta e prática
- Para executar ações, use as ferramentas disponíveis — não apenas descreva o que fazer
- Quando o usuário pede algo ambíguo, prefira agir e explicar brevemente o que fez
- Ao listar emails, seja conciso: assunto + remetente, sem repetir cabeçalhos
- Para pareceres tributários, sempre use a ferramenta gerar_parecer para iniciar o fluxo interativo"""

CLAUDE_TOOLS = [
    {
        "name": "ver_emails",
        "description": "Lista emails não lidos. Use para pedidos como 'o que chegou hoje?', 'tem algo urgente?', 'emails da semana', 'ver financeiro', etc.",
        "input_schema": {
            "type": "object",
            "properties": {
                "periodo": {
                    "type": "string",
                    "enum": ["hoje", "semana", "tudo"],
                    "description": "hoje=últimas 24h, semana=últimos 7d, tudo=caixa inteira"
                },
                "categoria": {
                    "type": "string",
                    "enum": ["URGENTE", "VENDAS", "PROFISSIONAL", "PESSOAL", "FINANCEIRO", "OUTROS"],
                    "description": "Filtrar por categoria (omitir para todas)"
                }
            }
        }
    },
    {
        "name": "contar_emails",
        "description": "Resumo geral com contagem de emails não lidos por categoria. Use para visão geral rápida ('quantos emails tenho?', 'resumo da caixa').",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "buscar_email",
        "description": "Busca o email mais recente de uma pessoa específica pelo nome.",
        "input_schema": {
            "type": "object",
            "properties": {
                "nome": {"type": "string", "description": "Nome ou parte do nome do remetente"}
            },
            "required": ["nome"]
        }
    },
    {
        "name": "analisar_email",
        "description": "Busca o email mais recente de um remetente e retorna análise com resumo e plano de ação.",
        "input_schema": {
            "type": "object",
            "properties": {
                "nome": {"type": "string", "description": "Nome do remetente"}
            },
            "required": ["nome"]
        }
    },
    {
        "name": "gerar_parecer",
        "description": "Inicia o fluxo interativo para gerar parecer técnico tributário. Use quando o usuário pedir 'parecer', 'análise tributária', 'orientação fiscal', etc.",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "listar_remetentes",
        "description": "Ranking de quem mais enviou emails não lidos. Use para 'quem mais manda email?', 'principais contatos', 'mapeamento de remetentes'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "periodo": {
                    "type": "string",
                    "enum": ["hoje", "semana", "tudo"],
                    "description": "Período de busca"
                }
            }
        }
    },
    {
        "name": "varredura_alertas",
        "description": "Varre toda a caixa procurando emails com palavras-chave críticas: clientes, SPED, PGFN, intimações, etc.",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "marcar_lidos",
        "description": "Marca todos os emails não lidos como lidos. Use para 'zera a caixa', 'marcar tudo como lido'.",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "reenviar_pdf",
        "description": "Reenvia o PDF do último parecer gerado como documento no Telegram, e/ou envia por email. Use quando o usuário pedir 'quero ver o pdf', 'me manda o parecer', 'envia o parecer', 'mostra o pdf'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "enviar_email": {
                    "type": "boolean",
                    "description": "Se true, também envia por email para o remetente original"
                }
            }
        }
    },
]

ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY) if ANTHROPIC_API_KEY else None

# ─── TRELLO ───────────────────────────────────────────────────────────────────
TRELLO_API_KEY = os.environ.get('TRELLO_API_KEY', '')
TRELLO_TOKEN = os.environ.get('TRELLO_TOKEN', '')
TRELLO_MESA1_LIST_ID = os.environ.get('TRELLO_MESA1_LIST_ID', '')
TRELLO_LABEL_URGENTE_ID = os.environ.get('TRELLO_LABEL_URGENTE_ID', '')

# ─── DRIVE PARECER ────────────────────────────────────────────────────────────
DRIVE_PARECER_ROOT_ID = os.environ.get('DRIVE_PARECER_ROOT_ID', '1Jd3oC3gNQoYXPUuKKBF8PrMZNUEkKT9m')

# ─── REGRAS AUTOMÁTICAS ───────────────────────────────────────────────────────
# Quando email bate numa regra: alerta Telegram + encaminha automaticamente

RULES = [
    {
        'name': 'IPOG — Solicitação de NF',
        'match_from': ['nf@ipog.edu.br', 'ipog.edu.br'],
        'match_subject': ['solicitação nf', 'solicitacao nf', 'nota fiscal', 'nf semana'],
        'forward_to': ['vanessa2mlima@gmail.com', 'vanessalima@compliance-ce.com.br'],
        'forward_from_account': 'cnp',
        'alert_msg': '🚨 IPOG — Solicitação de NF recebida!\nEncaminhado para Vanessa automaticamente.',
        'priority': 'URGENTE',
    },
    {
        'name': 'Sherman Alcantara — Fortes Tecnologia',
        'match_from': ['shermanalcantara@fortestecnologia.com.br'],
        'match_subject': [],   # qualquer assunto
        'match_logic': 'or',   # basta o remetente bater
        'action': 'trello_card',
        'trello_list_id': TRELLO_MESA1_LIST_ID,
        'alert_msg': '🚨 SHERMAN (Fortes) — Card criado na Mesa de Execução 1!',
        'priority': 'URGENTE',
    },
]

# Controle de IDs já processados por regras (evita repetição)
_rule_processed_ids = set()

# Controle de alertas já enviados (evita repetição)
_alerted_ids = set()

# Controle de solicitações de parecer pendentes: uid -> email_data
_parecer_pending = {}

# ─── ESTADO DO FLUXO INTERATIVO DE PARECER ───────────────────────────────────
_flow_counter = 0
_parecer_flow = {}        # fid -> state dict
_flow_folders_cache = {}  # fid -> [(name, id), ...]
_last_shown_email = None  # último email exibido na conversa (contexto)
_last_parecer = None      # último parecer gerado (pdf_bytes, urls, filename)

# ─── REDIS — persistência de estado ──────────────────────────────────────────

_redis_conn = None

def _redis():
    global _redis_conn
    if _redis_conn is None:
        url = os.environ.get('REDIS_URL', '')
        if url:
            try:
                _redis_conn = _redis_lib.from_url(url, decode_responses=False)
                _redis_conn.ping()
                print('Redis conectado.')
            except Exception as e:
                print(f'Redis indisponível: {e}')
    return _redis_conn

def _rget(key, default=None):
    r = _redis()
    if not r: return default
    try:
        v = r.get(f'bot:{key}')
        return json.loads(v) if v else default
    except: return default

def _rset(key, value, ttl=86400):
    r = _redis()
    if not r: return
    try: r.setex(f'bot:{key}', ttl, json.dumps(value, ensure_ascii=False, default=str))
    except: pass

def _rget_bytes(key):
    r = _redis()
    if not r: return None
    try: return r.get(f'bot:{key}')
    except: return None

def _rset_bytes(key, value, ttl=86400):
    r = _redis()
    if not r: return
    try: r.setex(f'bot:{key}', ttl, value)
    except: pass

# ─── Funções de persistência (write-through: global + Redis) ─────────────────

def persist_last_email(data):
    global _last_shown_email
    _last_shown_email = data
    _rset('last_email', data)

def persist_last_parecer(data):
    global _last_parecer
    _last_parecer = data
    pdf = data.get('pdf_bytes')
    docx = data.get('docx_bytes')
    meta = {k: v for k, v in data.items() if k not in ('pdf_bytes', 'docx_bytes')}
    _rset('last_parecer_meta', meta)
    if pdf:
        _rset_bytes('last_parecer_pdf', pdf)
    if docx:
        _rset_bytes('last_parecer_docx', docx)

def get_last_parecer():
    if _last_parecer:
        return _last_parecer
    meta = _rget('last_parecer_meta')
    if not meta:
        return None
    pdf = _rget_bytes('last_parecer_pdf')
    docx = _rget_bytes('last_parecer_docx')
    return {**meta, 'pdf_bytes': pdf, 'docx_bytes': docx}

def persist_parecer_pending():
    _rset('parecer_pending', _parecer_pending)

def load_state_from_redis():
    """Carrega estado persistido do Redis para a memória ao iniciar."""
    global _last_shown_email, _last_parecer, _parecer_pending, _alerted_ids, _rule_processed_ids
    email_data = _rget('last_email')
    if email_data:
        _last_shown_email = email_data
    meta = _rget('last_parecer_meta')
    if meta:
        pdf = _rget_bytes('last_parecer_pdf')
        _last_parecer = {**meta, 'pdf_bytes': pdf}
    pending = _rget('parecer_pending')
    if pending:
        _parecer_pending = pending
    alerted = _rget('alerted_ids')
    if alerted:
        _alerted_ids = set(alerted)
    rule_proc = _rget('rule_processed_ids')
    if rule_proc:
        _rule_processed_ids = set(rule_proc)
    print(f'Estado Redis carregado: {len(_parecer_pending)} pendentes, last_email={bool(_last_shown_email)}, last_parecer={bool(_last_parecer)}')

# ─── DETECÇÃO DE SOLICITAÇÕES DE PARECER ─────────────────────────────────────
# Palavras no ASSUNTO que sugerem consulta tributária/fiscal
PARECER_SUBJECT_KEYWORDS = [
    'fiscal', 'tributar', 'análise', 'analise', 'orientação', 'orientacao',
    'consulta', 'regime', 'cnae', 'icms', 'pis', 'cofins', 'posicionamento',
    'parecer', 'estrutura', 'enquadramento', 'filial', 'cnpj', 'nfc-e',
]

# Frases no CORPO que confirmam pedido de análise/orientação
PARECER_BODY_KEYWORDS = [
    'análise e orientação', 'analise e orientacao',
    'orientação sobre', 'orientacao sobre',
    'gostaríamos de receber', 'gostariamos de receber',
    'ficamos no seu aguardo', 'aguardamos seu retorno',
    'aguardo seu posicionamento', 'pedimos orientação', 'pedimos orientacao',
    'solicito orientação', 'solicito orientacao', 'solicito parecer',
    'posicionamento técnico', 'posicionamento tecnico',
    'necessidade de abertura', 'novo cnpj', 'nova filial',
    'adequação do regime', 'adequacao do regime',
    'regime tributário adequado', 'impacto tributário', 'impacto tributario',
    'detalhamento do projeto para análise', 'detalhamento do projeto para analise',
    'orientação técnica', 'orientacao tecnica',
    'parte fiscal', 'análise fiscal', 'analise fiscal',
    'estrutura tributária', 'estrutura tributaria',
    'icms-st', 'cnae adequado', 'obrigações tributárias', 'obrigacoes tributarias',
]


# ─── TRELLO HELPER ───────────────────────────────────────────────────────────

def create_trello_card_urgent(subject, sender, body_preview, list_id=None):
    """Cria card no Trello no topo da lista com label URGENTE."""
    if not TRELLO_API_KEY or not TRELLO_TOKEN:
        return None, 'TRELLO_API_KEY/TRELLO_TOKEN não configurados'
    lid = list_id or TRELLO_MESA1_LIST_ID
    if not lid:
        return None, 'TRELLO_MESA1_LIST_ID não configurado'

    name_display = sender.split('<')[0].strip().strip('"') or sender
    card_name = f'[FORTES] {subject[:90]}' if subject else f'[FORTES] Email de {name_display}'
    desc = f'**Remetente:** {sender}\n**Assunto:** {subject}\n\n{body_preview[:800]}'

    params = {
        'key': TRELLO_API_KEY,
        'token': TRELLO_TOKEN,
        'idList': lid,
        'name': card_name,
        'desc': desc,
        'pos': 'top',
    }
    if TRELLO_LABEL_URGENTE_ID:
        params['idLabels'] = TRELLO_LABEL_URGENTE_ID

    try:
        resp = requests.post('https://api.trello.com/1/cards', params=params, timeout=10)
        resp.raise_for_status()
        card = resp.json()
        return card.get('shortUrl'), None
    except Exception as e:
        return None, str(e)


# ─── HELPERS ─────────────────────────────────────────────────────────────────

def decode_str(s):
    if not s:
        return ''
    parts = decode_header(s)
    result = ''
    for part, enc in parts:
        if isinstance(part, bytes):
            result += part.decode(enc or 'utf-8', errors='replace')
        else:
            result += str(part)
    return result.strip()


def is_spam(sender, subject):
    s, sub = sender.lower(), subject.lower()
    return any(x in s for x in SPAM_SENDERS) or any(x in sub for x in SPAM_SUBJECTS)


def classify(sender, subject):
    s, sub = sender.lower(), subject.lower()
    if any(x in s for x in URGENT_SENDERS) or any(x in sub for x in URGENT_SUBJECTS):
        return 'URGENTE'
    if any(x in s for x in SALES_SENDERS):
        return 'VENDAS'
    if any(x in s for x in PROFESSIONAL_SENDERS):
        return 'PROFISSIONAL'
    if any(x in s for x in PERSONAL_SENDERS):
        return 'PESSOAL'
    if any(x in s for x in FINANCIAL_SENDERS):
        return 'FINANCEIRO'
    return 'OUTROS'


def match_alert(sender, subject):
    """Retorna o keyword que deu match, ou None."""
    combined = (sender + ' ' + subject).lower()
    for kw in ALERT_KEYWORDS:
        if kw in combined:
            return kw
    return None


def get_imap(user=None, password=None, readonly=True):
    user = user or GMAIL_USER
    password = password or GMAIL_APP_PASSWORD
    mail = imaplib.IMAP4_SSL('imap.gmail.com')
    mail.login(user, password)
    mail.select('inbox', readonly=readonly)
    return mail


def search_unseen(mail, hours=None):
    if hours:
        since = (datetime.now() - timedelta(hours=hours)).strftime('%d-%b-%Y')
        _, data = mail.search(None, f'(UNSEEN SINCE {since})')
    else:
        _, data = mail.search(None, 'UNSEEN')
    return data[0].split()


def fetch_headers(mail, ids, limit=200):
    results = []
    for eid in ids[-limit:]:
        _, msg_data = mail.fetch(eid, '(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT)])')
        raw = msg_data[0][1]
        msg = email.message_from_bytes(raw)
        sender = decode_str(msg.get('From', ''))
        subject = decode_str(msg.get('Subject', '(sem assunto)'))
        results.append((eid, sender, subject))
    return results


def match_rule(sender, subject):
    """Retorna a regra que bate com o email, ou None.
    match_logic 'and' (padrão): remetente E assunto devem bater.
    match_logic 'or': qualquer um dos dois basta.
    """
    s, sub = sender.lower(), subject.lower()
    for rule in RULES:
        from_match = any(x in s for x in rule['match_from'])
        subj_match = any(x in sub for x in rule['match_subject'])
        logic = rule.get('match_logic', 'and')
        matched = (from_match or subj_match) if logic == 'or' else (from_match and subj_match)
        if matched:
            return rule
    return None


def forward_email(original_msg_bytes, rule, acc):
    """Encaminha o email via SMTP usando a conta configurada na regra.
    Retorna (True, None) em caso de sucesso ou (False, mensagem_erro).
    """
    try:
        send_acc = next(
            (a for a in ACCOUNTS if rule['forward_from_account'] in a['user']),
            ACCOUNTS[0]
        )
        original = email.message_from_bytes(original_msg_bytes)
        orig_from = decode_str(original.get('From', ''))
        orig_subject = decode_str(original.get('Subject', ''))
        orig_date = decode_str(original.get('Date', ''))
        orig_body = get_email_body(original)

        fwd = MIMEMultipart()
        fwd['From'] = send_acc['user']
        fwd['To'] = ', '.join(rule['forward_to'])
        fwd['Subject'] = f'Fwd: {orig_subject}'

        body = (
            f'-------- Mensagem encaminhada --------\n'
            f'De: {orig_from}\n'
            f'Data: {orig_date}\n'
            f'Assunto: {orig_subject}\n\n'
            f'{orig_body}'
        )
        fwd.attach(MIMEText(body, 'plain', 'utf-8'))

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(send_acc['user'], send_acc['password'])
            smtp.sendmail(send_acc['user'], rule['forward_to'], fwd.as_bytes())
        return True, None
    except Exception as e:
        err = str(e)
        print(f'Erro ao encaminhar email [{send_acc["user"] if "send_acc" in dir() else "?"}]: {err}')
        return False, err


def get_email_body(msg):
    """Extrai o texto do corpo do email."""
    body = ''
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            cd = str(part.get('Content-Disposition', ''))
            if ct == 'text/plain' and 'attachment' not in cd:
                try:
                    body += part.get_payload(decode=True).decode(part.get_content_charset() or 'utf-8', errors='replace')
                except:
                    pass
    else:
        try:
            body = msg.get_payload(decode=True).decode(msg.get_content_charset() or 'utf-8', errors='replace')
        except:
            body = str(msg.get_payload())
    return body.strip()[:3000]  # limite para API


def search_email_by_name(name_query):
    """Busca email de um remetente priorizando não lidos, depois mais recente."""
    for acc in ACCOUNTS:
        try:
            mail = get_imap(acc['user'], acc['password'], readonly=True)

            def _fetch_result(eid):
                _, msg_data = mail.fetch(eid, '(BODY.PEEK[])')
                msg = email.message_from_bytes(msg_data[0][1])
                return {
                    'sender': decode_str(msg.get('From', '')),
                    'subject': decode_str(msg.get('Subject', '')),
                    'date': decode_str(msg.get('Date', '')),
                    'body': get_email_body(msg),
                    'account': acc['label'],
                }

            # 1. Tenta não lidos do remetente (prioridade máxima)
            _, data = mail.search(None, f'(UNSEEN FROM "{name_query}")')
            unseen_ids = data[0].split()
            if unseen_ids:
                result = _fetch_result(unseen_ids[-1])
                mail.logout()
                return result

            # 2. Busca qualquer email do remetente (lido ou não)
            _, data = mail.search(None, f'(FROM "{name_query}")')
            ids = data[0].split()

            # 3. Busca parcial nos últimos 200 se não achou pelo FROM direto
            if not ids:
                _, data = mail.search(None, 'ALL')
                all_ids = data[0].split()

                # Sub-busca: não lidos primeiro
                _, unseen_data = mail.search(None, 'UNSEEN')
                unseen_set = set(unseen_data[0].split())

                candidates = []
                for eid in reversed(all_ids[-200:]):
                    _, hdr = mail.fetch(eid, '(BODY.PEEK[HEADER.FIELDS (FROM)])')
                    sender = decode_str(email.message_from_bytes(hdr[0][1]).get('From', ''))
                    if name_query.lower() in sender.lower():
                        candidates.append((eid, eid in unseen_set))

                if candidates:
                    # Prefere não lido; se todos lidos, pega o primeiro (mais recente)
                    unread = [eid for eid, is_unseen in candidates if is_unseen]
                    chosen = unread[0] if unread else candidates[0][0]
                    result = _fetch_result(chosen)
                    mail.logout()
                    return result

            if ids:
                result = _fetch_result(ids[-1])
                mail.logout()
                return result

            mail.logout()
        except Exception as e:
            print(f'Erro search_email_by_name: {e}')
    return None


def match_parecer_subject(subject):
    """Retorna True se o assunto sugere uma consulta tributária."""
    sub = subject.lower()
    return any(kw in sub for kw in PARECER_SUBJECT_KEYWORDS)


def match_parecer_body(body):
    """Retorna True se o corpo confirma pedido de orientação/análise."""
    b = body.lower()
    return any(kw in b for kw in PARECER_BODY_KEYWORDS)


def extract_empresa_cnpj(text):
    """Extrai nome da empresa e CNPJ do corpo do email."""
    empresa = None
    cnpj = None

    m = re.search(r'\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}', text)
    if m:
        cnpj = m.group()

    patterns = [
        r'empresa\s+([A-ZÀ-Úa-zà-ú][^\(,\n]{2,40}?)(?:\s*[\(,]|\s+CNPJ|\s+\(CNPJ)',
        r'([A-ZÀ-Ú][A-ZÀ-Úa-zà-ú\s]{2,30})\s*\(CNPJ',
        r'cliente\s+([A-ZÀ-Úa-zà-ú][^\(,\n]{2,40}?)(?:\s*[\(,])',
    ]
    for pat in patterns:
        m2 = re.search(pat, text, re.IGNORECASE)
        if m2:
            empresa = m2.group(1).strip().rstrip('.,')
            break

    return empresa, cnpj


def generate_parecer_claude(email_data):
    """Gera parecer preliminar via Claude Sonnet seguindo as 10 seções da Compliance."""
    if not claude_client:
        return None
    empresa = email_data.get('empresa', 'Cliente')
    cnpj = email_data.get('cnpj', '')
    hoje = datetime.now().strftime('%d/%m/%Y')

    prompt = f"""Você é Marcos Lima, contador tributário da Compliance-CE.

Analise o email abaixo e elabore um PARECER TÉCNICO PRELIMINAR seguindo exatamente as 10 seções:

1. CABEÇALHO
   Data: {hoje} | Empresa: {empresa} | CNPJ: {cnpj}
   Solicitante: {email_data['sender']} | Parecerista: Marcos Lima

2. QUESTIONAMENTO
   Contexte a consulta, as dúvidas levantadas e a base legal citada pelo cliente.

3. BASE LEGAL ATUALIZADA
   Liste artigos, incisos e alíneas aplicáveis. Formato: "LC 214/2025, art. 27, § 3º, inciso II, alínea a"

4. INTERPRETAÇÃO NORMATIVA
   Análise técnica da norma, hermenêutica jurídico-contábil, posicionamento doutrinário quando relevante.

5. IMPLICAÇÕES PRÁTICAS
   Consequências concretas para o contribuinte: impacto fiscal, operacional e contábil.

6. ORIENTAÇÕES OPERACIONAIS
   O que o cliente deve fazer: prazos, procedimentos, ajustes em sistemas e documentos.

7. VEDAÇÕES OU EXCEÇÕES
   Restrições legais, exceções à regra geral, riscos de autuação fiscal.

8. CITAÇÕES COMPLEMENTARES
   Jurisprudência (STJ, TRFs, CARF), Soluções de Consulta RFB, manifestações SEFAZ, doutrina.

9. CONCLUSÃO
   Síntese objetiva — resposta direta a cada pergunta do cliente.

10. ASSINATURA TÉCNICA
    Marcos Lima — CRC/CE | Data: {hoje}
    ⚠️ PARECER PRELIMINAR — sujeito a revisão antes da entrega ao cliente.

---
Email recebido:
De: {email_data['sender']}
Assunto: {email_data['subject']}

{email_data['body']}
---

Responda em português, com linguagem técnica tributária. Cite a legislação de forma completa e precisa."""

    msg = claude_client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=4000,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return msg.content[0].text


def save_to_drive(content, empresa, cnpj=None):
    """Salva o parecer no Google Drive em Empresas/{empresa}/. Requer GOOGLE_SERVICE_ACCOUNT_JSON e DRIVE_EMPRESAS_FOLDER_ID."""
    sa_json = os.environ.get('GOOGLE_SERVICE_ACCOUNT_JSON', '')
    folder_id = os.environ.get('DRIVE_EMPRESAS_FOLDER_ID', '')
    if not sa_json or not folder_id:
        return None
    try:
        import json as _json
        import io as _io
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseUpload

        creds = service_account.Credentials.from_service_account_info(
            _json.loads(sa_json),
            scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=creds)

        empresa_clean = re.sub(r'[^\w\s\-]', '', empresa).strip()[:50]

        results = service.files().list(
            q=f"name='{empresa_clean}' and '{folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields='files(id)'
        ).execute()
        files = results.get('files', [])
        if files:
            subfolder_id = files[0]['id']
        else:
            meta = {'name': empresa_clean, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [folder_id]}
            f = service.files().create(body=meta, fields='id').execute()
            subfolder_id = f['id']

        ts = datetime.now().strftime('%Y%m%d_%H%M')
        filename = f'PARECER_PRELIMINAR_{empresa_clean}_{ts}.txt'
        media = MediaIoBaseUpload(_io.BytesIO(content.encode('utf-8')), mimetype='text/plain')
        file_meta = {'name': filename, 'parents': [subfolder_id]}
        created = service.files().create(body=file_meta, media_body=media, fields='id,webViewLink').execute()
        return created.get('webViewLink')
    except Exception as e:
        print(f'Erro save_to_drive: {e}')
        return None


def analyze_with_claude(email_data):
    """Envia email para Claude e retorna resumo + plano de ação."""
    if not claude_client:
        return '❌ API Claude não configurada.'
    prompt = f"""Você é um assistente de um contador tributário brasileiro chamado Marcos Lima.

Analise o email abaixo e retorne:

1. RESUMO (máx 3 linhas): do que se trata o email
2. PLANO DE AÇÃO: lista numerada de ações concretas que Marcos deve tomar, em ordem de prioridade
3. PRAZO SUGERIDO: quando resolver (se identificável)
4. QUEM ENVOLVER: se precisar de mais alguém da equipe

Email:
De: {email_data['sender']}
Assunto: {email_data['subject']}
Data: {email_data['date']}

{email_data['body']}

Responda em português, de forma direta e prática. Sem introduções."""

    msg = claude_client.messages.create(
        model='claude-haiku-4-5-20251001',
        max_tokens=600,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return msg.content[0].text


def fetch_emails_account(acc, hours=None, only_cat=None, limit=100):
    mail = get_imap(acc['user'], acc['password'], readonly=True)
    ids = search_unseen(mail, hours)
    total_unseen = len(ids)
    buckets = {k: [] for k in ICONS}
    spam_count = 0
    for eid, sender, subject in fetch_headers(mail, ids, limit):
        if is_spam(sender, subject):
            spam_count += 1
            continue
        cat = classify(sender, subject)
        if only_cat and cat != only_cat:
            continue
        name = sender.split('<')[0].strip().strip('"') or sender
        name = name[:35] + '...' if len(name) > 35 else name
        sub = subject[:55] + '...' if len(subject) > 55 else subject
        buckets[cat].append(f'• {sub}\n  ↳ {name}')
    mail.logout()
    return buckets, total_unseen, spam_count


def fetch_emails(hours=None, only_cat=None, limit=100):
    try:
        merged = {k: [] for k in ICONS}
        total_all, spam_all = 0, 0
        for acc in ACCOUNTS:
            buckets, total, spam = fetch_emails_account(acc, hours, only_cat, limit)
            total_all += total
            spam_all += spam
            for cat in ICONS:
                # Prefixar com label da conta se mais de uma conta
                if len(ACCOUNTS) > 1:
                    merged[cat].extend([f'[{acc["label"]}] {item}' for item in buckets[cat]])
                else:
                    merged[cat].extend(buckets[cat])
        return merged, total_all, spam_all, None
    except Exception as e:
        return None, 0, 0, str(e)


def fetch_by_sender(hours=None, limit=300):
    """Retorna lista de (nome, email, count, categoria) dos remetentes com mais emails não lidos."""
    try:
        from collections import defaultdict, Counter
        sender_counts = Counter()
        sender_meta = {}  # email_addr -> (nome_display, categoria)
        spam_count = 0
        for acc in ACCOUNTS:
            mail = get_imap(acc['user'], acc['password'], readonly=True)
            ids = search_unseen(mail, hours)
            for eid, sender, subject in fetch_headers(mail, ids, limit):
                if is_spam(sender, subject):
                    spam_count += 1
                    continue
                # Extrai nome e endereço
                m = re.match(r'^(.*?)\s*<([^>]+)>', sender)
                if m:
                    nome = m.group(1).strip().strip('"') or m.group(2)
                    addr = m.group(2).lower()
                else:
                    nome = sender.strip()
                    addr = sender.lower()
                cat = classify(sender, subject)
                sender_counts[addr] += 1
                if addr not in sender_meta:
                    sender_meta[addr] = (nome[:40], cat)
            mail.logout()
        ranked = sorted(sender_counts.items(), key=lambda x: x[1], reverse=True)
        result = [(sender_meta[addr][0], addr, count, sender_meta[addr][1]) for addr, count in ranked]
        return result, spam_count, None
    except Exception as e:
        return [], 0, str(e)


def count_by_category(hours=None, limit=500):
    try:
        total_all = 0
        counts_all = {k: 0 for k in ICONS}
        spam_all = 0
        for acc in ACCOUNTS:
            mail = get_imap(acc['user'], acc['password'], readonly=True)
            ids = search_unseen(mail, hours)
            total_all += len(ids)
            for eid, sender, subject in fetch_headers(mail, ids, limit):
                if is_spam(sender, subject):
                    spam_all += 1
                    continue
                counts_all[classify(sender, subject)] += 1
            mail.logout()
        return total_all, counts_all, spam_all, None
    except Exception as e:
        return 0, {}, 0, str(e)


def mark_as_read(hours=None):
    try:
        total = 0
        for acc in ACCOUNTS:
            mail = get_imap(acc['user'], acc['password'], readonly=False)
            ids = search_unseen(mail, hours)
            if ids:
                for i in range(0, len(ids), 100):
                    batch = ids[i:i+100]
                    mail.store(b','.join(batch), '+FLAGS', '\\Seen')
            total += len(ids)
            mail.logout()
        return total
    except Exception as e:
        return -1


def build_message(buckets, title='📬 Emails não lidos'):
    lines = [title, '']
    total = 0
    for cat, items in buckets.items():
        if items:
            lines.append(f'{ICONS[cat]} {cat}')
            lines.extend(items)
            lines.append('')
            total += len(items)
    if total == 0:
        return '✅ Nenhum email importante no momento.'
    return '\n'.join(lines).strip()


# ─── DRIVE PARECER — HELPERS ─────────────────────────────────────────────────

def _drive_service():
    sa_json = os.environ.get('GOOGLE_SERVICE_ACCOUNT_JSON', '')
    if not sa_json:
        return None
    import json as _j
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    creds = service_account.Credentials.from_service_account_info(
        _j.loads(sa_json), scopes=['https://www.googleapis.com/auth/drive']
    )
    return build('drive', 'v3', credentials=creds)


def list_drive_subfolders(parent_id):
    """Lista TODAS as subpastas usando paginação do Drive API."""
    svc = _drive_service()
    if not svc:
        return []
    try:
        folders = []
        page_token = None
        while True:
            kwargs = dict(
                q=f"'{parent_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false",
                fields='nextPageToken, files(id,name)',
                orderBy='name',
                pageSize=100,
            )
            if page_token:
                kwargs['pageToken'] = page_token
            res = svc.files().list(**kwargs).execute()
            folders.extend([(f['name'], f['id']) for f in res.get('files', [])])
            page_token = res.get('nextPageToken')
            if not page_token:
                break
        return folders
    except Exception as e:
        print(f'list_drive_subfolders: {e}')
        return []


def create_drive_folder(name, parent_id):
    svc = _drive_service()
    if not svc:
        return None, None
    try:
        meta = {'name': name, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [parent_id]}
        f = svc.files().create(body=meta, fields='id,webViewLink').execute()
        return f['id'], f.get('webViewLink', '')
    except Exception as e:
        print(f'create_drive_folder: {e}')
        return None, None


def upload_to_drive_as_gdoc(docx_bytes, base_filename, folder_id):
    """Upload DOCX as Google Doc format (no service-account storage quota consumed).
    Returns (gdoc_url, pdf_bytes, err_msg).
    Google-native files don't count against service account quota — only binary files do."""
    svc = _drive_service()
    if not svc:
        return None, None, 'GOOGLE_SERVICE_ACCOUNT_JSON nao configurado'
    from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
    DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    gdoc_url = None
    pdf_bytes = None
    err = None
    # Step 1: Upload DOCX converting to Google Docs format (free of binary quota)
    try:
        gdoc_meta = {
            'name': base_filename,
            'parents': [folder_id],
            'mimeType': 'application/vnd.google-apps.document',
        }
        gdoc = svc.files().create(
            body=gdoc_meta,
            media_body=MediaIoBaseUpload(_io_module.BytesIO(docx_bytes), mimetype=DOCX_MIME),
            fields='id,webViewLink'
        ).execute()
        gdoc_url = gdoc.get('webViewLink', '')
        gdoc_id = gdoc['id']
    except Exception as e:
        err = f'gdoc_upload: {e}'
        print(err)
        return None, None, err
    # Step 2: Export Google Doc as PDF bytes (read operation, no quota)
    try:
        req = svc.files().export_media(fileId=gdoc_id, mimeType='application/pdf')
        pdf_buf = _io_module.BytesIO()
        dl = MediaIoBaseDownload(pdf_buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        pdf_bytes = pdf_buf.getvalue()
    except Exception as e:
        err = f'pdf_export: {e}'
        print(err)
    return gdoc_url, pdf_bytes, err


def generate_parecer_docx(parecer_text, empresa, cnpj, modelo):
    """Returns BytesIO with formatted DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    titulo_doc = 'NOTA TÉCNICA TRIBUTÁRIA' if modelo == 'EMPRESARIAL' else 'PARECER JURÍDICO - TRIBUTÁRIO E FISCAL'
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(titulo_doc)
    r.bold = True
    r.font.size = Pt(13)
    t.paragraph_format.space_after = Pt(4)

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(empresa + (f'  |  CNPJ: {cnpj}' if cnpj else '') + '  |  Compliance-CE')
    c.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    def _is_secao(line):
        return bool(re.match(
            r'^(CABECALHO|EMENTA|ASSINATURA|I\.|II\.|III\.|IV\.|II\.\d|III\.\d|IV\.\d)',
            line.strip()
        ))

    for line in parecer_text.split('\n'):
        line_str = line.strip()
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = Pt(18)
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)

        if not line_str:
            pf.space_after = Pt(2)
            continue

        if _is_secao(line_str):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Inches(0)
            pf.space_before = Pt(14)
            r = p.add_run(line_str)
            r.bold = True
            r.font.size = Pt(11)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Inches(0.79)  # 2 cm ≈ 0.79 inch
            r = p.add_run(line_str)
            r.font.size = Pt(11)

    buf = _io_module.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _sanitize_pdf_text(text):
    """Replace characters outside CP1252 so fpdf2 core fonts don't fail."""
    replacements = {
        '—': '-', '–': '-',   # em dash, en dash
        '‘': "'", '’': "'",   # curly single quotes
        '“': '"', '”': '"',   # curly double quotes
        '…': '...', ' ': ' ', # ellipsis, non-breaking space
        '•': '-', '‣': '-',   # bullet points
        '●': '-', '►': '>',
        '━': '-', '│': '|',   # box-drawing chars
        '−': '-', '·': '-',   # minus sign, middle dot
        '×': 'x', '÷': '/',  # × ÷
        '―': '-',                  # horizontal bar
        '▪': '-', '■': '-',
    }
    for ch, rep in replacements.items():
        text = text.replace(ch, rep)
    # Final safety: encode to CP1252, replace any remaining unrepresentable chars
    return text.encode('cp1252', errors='replace').decode('cp1252')


def generate_parecer_pdf_bytes(parecer_text, empresa, cnpj, modelo):
    """Generate PDF directly via fpdf2 (no Google Docs dependency)."""
    try:
        from fpdf import FPDF
    except ImportError:
        print('fpdf2 not installed')
        return None
    try:
        titulo = 'PARECER JURIDICO - TRIBUTARIO E FISCAL' if modelo == 'GERAL' else 'NOTA TECNICA TRIBUTARIA'
        empresa_s = _sanitize_pdf_text(empresa)
        cnpj_s = _sanitize_pdf_text(cnpj)
        texto_s = _sanitize_pdf_text(parecer_text)

        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_margins(left=25, top=25, right=20)
        pdf.add_page()

        # Title block
        pdf.set_font('Helvetica', 'B', 13)
        pdf.multi_cell(0, 7, titulo, align='C')
        subtitulo = empresa_s + (f'  |  CNPJ: {cnpj_s}' if cnpj_s else '') + '  |  Compliance-CE'
        pdf.set_font('Helvetica', '', 10)
        pdf.multi_cell(0, 6, subtitulo, align='C')
        pdf.ln(5)

        # Body
        for line in texto_s.split('\n'):
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue
            is_sec = bool(re.match(
                r'^(CABECALHO|EMENTA|ASSINATURA|I\.|II\.|III\.|IV\.|II\.\d|III\.\d|IV\.\d)',
                stripped
            ))
            if is_sec:
                pdf.ln(3)
                pdf.set_font('Helvetica', 'B', 11)
                pdf.multi_cell(0, 6, stripped, align='J')
                pdf.set_font('Helvetica', '', 11)
            else:
                pdf.set_font('Helvetica', '', 11)
                pdf.multi_cell(0, 6, stripped, align='J')
                pdf.ln(1)

        return bytes(pdf.output())
    except Exception as e:
        print(f'generate_parecer_pdf_bytes error: {e}')
        return None


def generate_parecer_claude_v2(email_data, modelo):
    """Generates parecer following MCP Compliance 9-section structure."""
    if not claude_client:
        return None
    empresa = email_data.get('empresa', 'Cliente')
    cnpj = email_data.get('cnpj', '')
    hoje = datetime.now().strftime('%d/%m/%Y')
    meses = ['janeiro','fevereiro','março','abril','maio','junho',
             'julho','agosto','setembro','outubro','novembro','dezembro']
    hoje_obj = datetime.now()
    hoje_extenso = f'{hoje_obj.day} de {meses[hoje_obj.month-1]} de {hoje_obj.year}'
    cnpj_str = f' | CNPJ: {cnpj}' if cnpj else ''

    if modelo == 'GERAL':
        titulo = 'PARECER JURIDICO - TRIBUTARIO E FISCAL'
        profundidade = (
            'Parecer tecnico completo. Maximo 5 paginas. '
            'Citar artigo, paragrafo, inciso, alinea em CADA afirmacao legal. '
            'Ex: "nos termos do art. 12, par. 2, inciso I, alinea b, da LC 214/2025". '
            'Todas as 9 secoes em extensao integral.'
        )
        base_legal_regra = 'Citar dispositivos legais completos no corpo de cada subsecao.'
    else:
        titulo = 'NOTA TECNICA TRIBUTARIA'
        profundidade = (
            'Nota tecnica condensada. Maximo 3 paginas. '
            'Linguagem direta para o empresario: "imposto" em vez de "tributo", '
            '"nota fiscal" em vez de "NF-e", "boleto" em vez de "DAE". '
            'Todas as 9 secoes obrigatorias, mas condensadas ao essencial.'
        )
        base_legal_regra = 'Base legal APENAS na tabela da secao II.7. NAO citar artigos no corpo do texto.'

    prompt = f"""Voce e Marcos Lima, contador tributario da Compliance-CE (Fortaleza/CE).
Elabore um {titulo} seguindo EXATAMENTE a estrutura abaixo.

MODO: {modelo} — {profundidade}

=== ESTRUTURA OBRIGATORIO (copiar os cabecalhos exatamente) ===

CABECALHO
CONSULENTE: {empresa}{cnpj_str}
DATA: {hoje}
ASSUNTO: {email_data['subject']}
LEGISLACAO: [listar as principais normas aplicadas, ex: LC 214/2025, CTN, RIR/2018]

EMENTA: [TODO EM MAIUSCULAS. 4 a 6 linhas. Tema + enquadramento legal + tese juridica + conclusao. Elaborar somente apos concluir a analise.]

I. RELATORIO
[Contextualizar a consulta: quem pergunta, o que pergunta, fatos relevantes. PROIBIDO antecipar conclusoes ou citar artigos nesta secao.]

II. FUNDAMENTACAO

II.1 Contexto do Questionamento
[Apresentar o contexto tecnico e relevancia da materia.]

II.2 Base Legal Atualizada
[{base_legal_regra} Incluir vigencia e eventuais alteracoes recentes.]

II.3 Interpretacao Normativa
[Analise tecnica: hermeneutica, ratio legis, posicao administrativa (RFB, SEFAZ-CE).]

II.4 Implicacoes Praticas
[Consequencias: fiscal, operacional, contabil. O que muda na rotina do cliente.]

II.5 Orientacoes Operacionais
[O que o cliente deve fazer: prazos, procedimentos, documentos necessarios.]

II.6 Vedacoes e Excecoes
[O que NAO pode ser feito. Riscos de autuacao. Situacoes de excecao.]

II.7 Citacoes Complementares
[Jurisprudencia do CARF, Solucoes de Consulta RFB, orientacoes SEFAZ. Citar numero/data.]

III. EXEMPLOS PRATICOS
[INCLUIR SOMENTE SE o tema envolver calculos tributarios, regimes comparativos, split payment, Simples Nacional. OMITIR se o tema for puramente interpretativo.]

IV. CONCLUSAO
1. [Resposta direta ao questionamento 1]
2. [Resposta direta ao questionamento 2, se houver]
[Numerada. Clara e executavel sem leitura previa da Fundamentacao.]

ASSINATURA
E o parecer, salvo melhor juizo.
Fortaleza, {hoje_extenso}.
Marcos Lima — Contador e Cientista de Dados | CRC/CE n. 23.224

=== REGRAS INVIOLAVEIS ===

1. TRAVESSAO (—) PROIBIDO EM QUALQUER PARTE DO TEXTO. Substitua por virgula ou dois-pontos.
2. PALAVRAS PROIBIDAS: "cumpre salientar", "destarte", "outrossim", "no que tange", "consoante",
   "em que pese", "assim sendo", "nesse diapasao", "verifica-se que", "nota-se que",
   "observa-se que", "ressalte-se que", "importa destacar", "urge destacar", "mister se faz",
   "afigura-se", "ab initio", "ex vi", "precipuamente", "sobremodo", "nesse passo".
   Substitua por: "portanto", "alem disso", "conforme", "quanto a", "apesar de", "assim".
3. Um paragrafo = uma ideia. Maximo 4 paragrafos por subsecao.
4. Frases curtas. Maximo 2 oracoes subordinadas por frase.
5. Secoes numeradas com algarismo romano (I, II, III). Subsecoes: II.1, II.2, etc.
6. Linguagem tecnica formal e direta. Sem adjetivos desnecessarios.

=== EMAIL RECEBIDO ===

De: {email_data['sender']}
Assunto: {email_data['subject']}

{email_data['body'][:3000]}

Responda SOMENTE com o texto do parecer no formato indicado. Sem preambulo, sem meta-comentarios."""

    msg = claude_client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=6000,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return msg.content[0].text


def reply_email_with_pdf(email_data, pdf_bytes, base_filename, summary):
    """Sends a reply email with the PDF attached."""
    try:
        acc = ACCOUNTS[0]
        for a in ACCOUNTS:
            if email_data.get('account', '') == a['label']:
                acc = a
                break

        original_sender = email_data.get('sender', '')
        m = re.match(r'.*<([^>]+)>', original_sender)
        reply_to = m.group(1) if m else original_sender

        msg = MIMEMultipart()
        msg['From'] = acc['user']
        msg['To'] = reply_to
        msg['Subject'] = f'Re: {email_data["subject"]}'
        msg.attach(MIMEText(summary, 'plain', 'utf-8'))

        part = MIMEBase('application', 'pdf')
        part.set_payload(pdf_bytes)
        _email_encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{base_filename}.pdf"')
        msg.attach(part)

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(acc['user'], acc['password'])
            smtp.sendmail(acc['user'], [reply_to], msg.as_bytes())
        return True, None
    except Exception as e:
        return False, str(e)


# ─── MONITORAMENTO AUTOMÁTICO (a cada 10 min) ─────────────────────────────────

async def check_alerts(context):
    """Roda em background — verifica alertas e regras automáticas."""
    global _alerted_ids, _rule_processed_ids
    for acc in ACCOUNTS:
        try:
            mail = get_imap(acc['user'], acc['password'], readonly=True)
            since = (datetime.now() - timedelta(minutes=35)).strftime('%d-%b-%Y')
            _, data = mail.search(None, f'(UNSEEN SINCE {since})')
            ids = data[0].split()

            for eid in ids[-50:]:
                uid = f"{acc['user']}:{eid}"

                # Busca headers
                _, msg_data = mail.fetch(eid, '(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT)])')
                raw = msg_data[0][1]
                msg = email.message_from_bytes(raw)
                sender = decode_str(msg.get('From', ''))
                subject = decode_str(msg.get('Subject', ''))

                # ── Verificar REGRAS automáticas ──
                rule = match_rule(sender, subject)
                if rule and uid not in _rule_processed_ids:
                    _rule_processed_ids.add(uid)
                    _rset('rule_processed_ids', list(_rule_processed_ids))
                    _, full_data = mail.fetch(eid, '(BODY.PEEK[])')
                    full_raw = full_data[0][1]
                    hora = datetime.now().strftime('%H:%M')

                    action = rule.get('action', 'forward')

                    if action == 'trello_card':
                        full_msg = email.message_from_bytes(full_raw)
                        body = get_email_body(full_msg)
                        lid = rule.get('trello_list_id') or TRELLO_MESA1_LIST_ID
                        card_url, err_msg = create_trello_card_urgent(subject, sender, body, list_id=lid)
                        if card_url:
                            detail = f'✅ Card criado no topo (URGENTE):\n{card_url}'
                        else:
                            detail = f'❌ Falha ao criar card Trello\n⚠️ {(err_msg or "erro desconhecido")[:300]}'
                    else:
                        ok, err_msg = forward_email(full_raw, rule, acc)
                        destinos = '\n'.join(f'  • {d}' for d in rule['forward_to'])
                        if ok:
                            detail = f'✅ Encaminhado para:\n{destinos}'
                        else:
                            detail = f'❌ Falha ao encaminhar para:\n{destinos}\n\n⚠️ Erro: {(err_msg or "desconhecido")[:300]}'

                    msg_tg = (
                        f'{rule["alert_msg"]}\n\n'
                        f'📧 {subject}\n'
                        f'↳ {sender}\n'
                        f'🕐 {hora}\n\n'
                        f'{detail}'
                    )
                    await context.bot.send_message(chat_id=CHAT_ID, text=msg_tg)
                    continue  # já tratado pela regra

                # ── Verificar SOLICITAÇÕES DE PARECER ──
                if uid not in _alerted_ids and match_parecer_subject(subject):
                    # Busca corpo para confirmação e extração de empresa
                    try:
                        _, full_data = mail.fetch(eid, '(BODY.PEEK[])')
                        full_raw = full_data[0][1]
                        full_msg = email.message_from_bytes(full_raw)
                        body = get_email_body(full_msg)
                        date_str = decode_str(full_msg.get('Date', ''))

                        if match_parecer_body(body):
                            _alerted_ids.add(uid)
                            empresa, cnpj = extract_empresa_cnpj(body)
                            name = sender.split('<')[0].strip().strip('"') or sender
                            hora = datetime.now().strftime('%H:%M')

                            _parecer_pending[uid] = {
                                'sender': sender,
                                'subject': subject,
                                'body': body,
                                'date': date_str,
                                'empresa': empresa or 'Cliente',
                                'cnpj': cnpj or '',
                                'account': acc['label'],
                            }
                            persist_parecer_pending()

                            empresa_info = f'\n🏢 Empresa: {empresa}' if empresa else ''
                            cnpj_info = f'\n📋 CNPJ: {cnpj}' if cnpj else ''
                            msg_tg = (
                                f'📝 SOLICITAÇÃO DE PARECER — {hora}\n'
                                f'{acc["label"]}{empresa_info}{cnpj_info}\n\n'
                                f'📧 {subject}\n'
                                f'↳ {name}\n\n'
                                f'💡 Responda "parecer" para gerar o parecer preliminar'
                            )
                            await context.bot.send_message(chat_id=CHAT_ID, text=msg_tg)
                            continue
                    except Exception as e:
                        print(f'Erro ao verificar parecer {uid}: {e}')

                # ── Verificar ALERTAS por palavra-chave ──
                if uid not in _alerted_ids:
                    kw = match_alert(sender, subject)
                    if kw:
                        _alerted_ids.add(uid)
                        _rset('alerted_ids', list(_alerted_ids))
                        name = sender.split('<')[0].strip().strip('"') or sender
                        hora = datetime.now().strftime('%H:%M')
                        msg_tg = (
                            f'🚨 ALERTA — {hora}\n'
                            f'{acc["label"]}\n\n'
                            f'Palavra-chave: {kw.upper()}\n\n'
                            f'📧 {subject}\n'
                            f'↳ {name}'
                        )
                        await context.bot.send_message(chat_id=CHAT_ID, text=msg_tg)

            mail.logout()
        except Exception as e:
            print(f'Erro check_alerts {acc["user"]}: {e}')


# ─── HANDLERS ────────────────────────────────────────────────────────────────

def only_authorized(func):
    async def wrapper(update: Update, context: ContextTypes.DEFAULT_TYPE):
        if update.effective_chat.id != CHAT_ID:
            return
        await func(update, context)
    return wrapper


@only_authorized
async def cmd_todos(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Contando TODOS os emails não lidos...')
    total, counts, spam, err = count_by_category(hours=None, limit=500)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    lines = [f'📊 RESUMO GERAL — {total} emails não lidos\n']
    for cat, n in counts.items():
        if n > 0:
            lines.append(f'{ICONS[cat]} {cat}: {n}')
    lines.append(f'🗑️ Propaganda: {spam}')
    nao_contados = total - sum(counts.values()) - spam
    if nao_contados > 0:
        lines.append(f'⚠️ Além do limite analisado: ~{nao_contados}')
    lines.append('\nUse "urgente", "vendas" ou "hoje" para detalhes.')
    await update.message.reply_text('\n'.join(lines))


@only_authorized
async def cmd_emails(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Emails das últimas 24h...')
    buckets, total, spam, err = fetch_emails(hours=24, limit=100)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    await update.message.reply_text(build_message(buckets, f'📬 Hoje — {total} total, {spam} propagandas'))


@only_authorized
async def cmd_hoje(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await cmd_emails(update, context)


@only_authorized
async def cmd_semana(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Emails da semana (7 dias)...')
    buckets, total, spam, err = fetch_emails(hours=168, limit=200)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    await update.message.reply_text(build_message(buckets, f'📬 Semana — {total} total, {spam} propagandas'))


@only_authorized
async def cmd_urgente(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Buscando urgentes...')
    buckets, total, spam, err = fetch_emails(hours=None, only_cat='URGENTE', limit=200)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    items = buckets.get('URGENTE', [])
    if not items:
        await update.message.reply_text('✅ Nenhum urgente.')
        return
    msg = f'🔴 URGENTE ({len(items)})\n\n' + '\n'.join(items[:30])
    await update.message.reply_text(msg)
    if len(items) > 30:
        await update.message.reply_text(f'... e mais {len(items)-30} emails urgentes.')


@only_authorized
async def cmd_vendas(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Buscando vendas...')
    buckets, total, spam, err = fetch_emails(hours=None, only_cat='VENDAS', limit=200)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    items = buckets.get('VENDAS', [])
    msg = (f'💰 VENDAS ({len(items)})\n\n' + '\n'.join(items[:30])) if items else '📭 Nenhuma venda.'
    await update.message.reply_text(msg)


@only_authorized
async def cmd_profissional(update: Update, context: ContextTypes.DEFAULT_TYPE):
    buckets, total, spam, err = fetch_emails(hours=None, only_cat='PROFISSIONAL', limit=100)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    items = buckets.get('PROFISSIONAL', [])
    msg = (f'👨‍💼 PROFISSIONAL ({len(items)})\n\n' + '\n'.join(items[:30])) if items else '📭 Nenhum.'
    await update.message.reply_text(msg)


@only_authorized
async def cmd_financeiro(update: Update, context: ContextTypes.DEFAULT_TYPE):
    buckets, total, spam, err = fetch_emails(hours=None, only_cat='FINANCEIRO', limit=100)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    items = buckets.get('FINANCEIRO', [])
    msg = (f'💳 FINANCEIRO ({len(items)})\n\n' + '\n'.join(items[:30])) if items else '📭 Nenhum email financeiro.'
    await update.message.reply_text(msg)


@only_authorized
async def cmd_pessoal(update: Update, context: ContextTypes.DEFAULT_TYPE):
    buckets, total, spam, err = fetch_emails(hours=None, only_cat='PESSOAL', limit=100)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    items = buckets.get('PESSOAL', [])
    msg = (f'👦 PESSOAL ({len(items)})\n\n' + '\n'.join(items[:30])) if items else '📭 Nenhum email pessoal.'
    await update.message.reply_text(msg)


@only_authorized
async def cmd_marcar_lido(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('⏳ Marcando TODOS como lidos...')
    total = mark_as_read(hours=None)
    if total == -1:
        await update.message.reply_text('❌ Erro ao marcar emails.')
    elif total == 0:
        await update.message.reply_text('📭 Nenhum email não lido.')
    else:
        await update.message.reply_text(f'✅ {total} email(s) marcados como lidos!')


@only_authorized
async def cmd_analisar_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Busca email por nome e retorna análise + plano de ação."""
    text = update.message.text.strip()
    # Extrai o nome após palavras-chave
    for kw in ['email do', 'email da', 'email de', 'analisar email do', 'analisar email da',
                'analisar email de', 'plano de acao', 'plano de ação', 'o que fazer com']:
        if kw in text.lower():
            name_query = text.lower().replace(kw, '').strip().title()
            break
    else:
        name_query = text.strip()

    if not name_query or len(name_query) < 2:
        await update.message.reply_text('❓ Qual o nome? Ex: "email da Raquel"')
        return

    await update.message.reply_text(f'🔍 Buscando email de {name_query}...')
    email_data = search_email_by_name(name_query)

    if not email_data:
        await update.message.reply_text(f'📭 Nenhum email encontrado de "{name_query}".')
        return

    await update.message.reply_text(
        f'📧 Encontrado!\nDe: {email_data["sender"]}\nAssunto: {email_data["subject"]}\n\n⏳ Analisando com IA...'
    )
    analysis = analyze_with_claude(email_data)
    await update.message.reply_text(analysis)


def scan_inbox_for_parecer():
    """Varre emails NÃO LIDOS buscando solicitações de parecer ainda não detectadas.
    Adiciona novos candidatos a _parecer_pending. Retorna quantidade encontrada.
    """
    found = 0
    for acc in ACCOUNTS:
        try:
            mail = get_imap(acc['user'], acc['password'], readonly=True)
            _, data = mail.search(None, 'UNSEEN')
            ids = data[0].split()
            for eid in ids[-100:]:
                uid = f"{acc['user']}:{eid}"
                if uid in _parecer_pending or uid in _alerted_ids:
                    continue
                _, msg_data = mail.fetch(eid, '(BODY.PEEK[HEADER.FIELDS (FROM SUBJECT)])')
                raw = msg_data[0][1]
                msg = email.message_from_bytes(raw)
                sender = decode_str(msg.get('From', ''))
                subject = decode_str(msg.get('Subject', ''))
                if not match_parecer_subject(subject):
                    continue
                _, full_data = mail.fetch(eid, '(BODY.PEEK[])')
                full_msg = email.message_from_bytes(full_data[0][1])
                body = get_email_body(full_msg)
                if not match_parecer_body(body):
                    continue
                empresa, cnpj = extract_empresa_cnpj(body)
                _parecer_pending[uid] = {
                    'sender': sender,
                    'subject': subject,
                    'body': body,
                    'date': decode_str(full_msg.get('Date', '')),
                    'empresa': empresa or 'Cliente',
                    'cnpj': cnpj or '',
                    'account': acc['label'],
                }
                _alerted_ids.add(uid)
                found += 1
            persist_parecer_pending()
            _rset('alerted_ids', list(_alerted_ids))
            mail.logout()
        except Exception as e:
            print(f'scan_inbox_for_parecer [{acc["user"]}]: {e}')
    return found


@only_authorized
async def cmd_gerar_parecer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Inicia o fluxo interativo de geração de parecer."""
    global _flow_counter

    if not _parecer_pending:
        await update.message.reply_text('🔍 Varrendo emails não lidos em busca de solicitações...')
        found = scan_inbox_for_parecer()
        if not found:
            # Usa o email exibido na conversa (memória Redis) como base
            ctx_email = _last_shown_email or _rget('last_email')
            if ctx_email:
                uid = f"context:{hash(ctx_email.get('subject',''))}"
                _parecer_pending[uid] = {
                    **ctx_email,
                    'empresa': ctx_email.get('empresa') or ctx_email['sender'].split('<')[0].strip().strip('"'),
                    'cnpj': ctx_email.get('cnpj', ''),
                }
                persist_parecer_pending()
                await update.message.reply_text(
                    f'📧 Usando o email da memória:\n'
                    f'📌 {ctx_email["subject"]}\n'
                    f'↳ {ctx_email["sender"]}'
                )
            else:
                await update.message.reply_text(
                    '📭 Nenhum email encontrado para o parecer.\n\n'
                    'Primeiro mostre um email ("me mostre o email da Ingrid") e depois peça o parecer.'
                )
                return
        else:
            await update.message.reply_text(f'✅ {found} solicitação(ões) encontrada(s)!')

    uid, email_data = list(_parecer_pending.items())[-1]

    _flow_counter += 1
    fid = _flow_counter
    _parecer_flow[fid] = {'uid': uid, 'email_data': email_data, 'state': 'folder_select'}

    await update.message.reply_text('⏳ Carregando pastas do Drive...')

    folders = list_drive_subfolders(DRIVE_PARECER_ROOT_ID)
    _flow_folders_cache[fid] = folders

    from telegram import InlineKeyboardButton, InlineKeyboardMarkup
    buttons = [
        [InlineKeyboardButton(f'📁 {name[:45]}', callback_data=f'fs:{fid}:{idx}')]
        for idx, (name, _) in enumerate(folders)
    ]
    buttons.append([InlineKeyboardButton('➕ Criar nova pasta', callback_data=f'fn:{fid}')])

    empresa = email_data.get('empresa', 'Cliente')
    cnpj = email_data.get('cnpj', '')
    await update.message.reply_text(
        f'📂 Selecione a pasta de destino:\n🏢 {empresa}' + (f'  |  {cnpj}' if cnpj else ''),
        reply_markup=InlineKeyboardMarkup(buttons)
    )


async def handle_callback_query(update, context):
    from telegram import InlineKeyboardButton, InlineKeyboardMarkup
    query = update.callback_query
    await query.answer()

    if query.message.chat.id != CHAT_ID:
        return

    parts = query.data.split(':', 2)
    action = parts[0]
    try:
        fid = int(parts[1])
    except (IndexError, ValueError):
        return
    extra = parts[2] if len(parts) > 2 else None

    if fid not in _parecer_flow:
        await query.edit_message_text('❌ Sessão expirada. Responda "parecer" novamente.')
        return

    flow = _parecer_flow[fid]

    if action == 'fs':
        idx = int(extra)
        folders = _flow_folders_cache.get(fid, [])
        name, folder_id = folders[idx]
        flow['folder_id'] = folder_id
        flow['folder_name'] = name
        flow['state'] = 'modelo_select'
        buttons = [
            [InlineKeyboardButton('📋 MODELO GERAL', callback_data=f'mg:{fid}')],
            [InlineKeyboardButton('🏢 MODELO EMPRESARIAL', callback_data=f'me:{fid}')],
        ]
        await query.edit_message_text(
            f'✅ Pasta: {name}\n\n📝 Selecione o modelo do parecer:',
            reply_markup=InlineKeyboardMarkup(buttons)
        )

    elif action == 'fn':
        flow['state'] = 'folder_new'
        await query.edit_message_text(
            '📝 Digite o nome da nova pasta no Drive:\n'
            '(ex: "Empresa XYZ — Reforma Tributária")'
        )

    elif action in ('mg', 'me'):
        flow['modelo'] = 'GERAL' if action == 'mg' else 'EMPRESARIAL'
        flow['state'] = 'generating'
        await query.edit_message_text(f'⏳ Gerando parecer {flow["modelo"]}... aguarde até 1 minuto.')
        await _gerar_e_salvar(context, fid, flow)

    elif action == 'as':
        await _enviar_email_resposta(query, context, fid, flow)

    elif action == 'ac':
        gdoc_url = flow.get('docx_url', '')  # docx_url armazena o link do Google Doc
        _parecer_flow.pop(fid, None)
        if gdoc_url:
            lines = ['📄 Parecer salvo no Drive como Google Doc. Email não enviado.\n',
                     'Para enviar depois, responda "envia o parecer" ou "manda o pdf".\n',
                     f'📄 Google Doc: {gdoc_url}']
        else:
            lines = ['⚠️ Parecer não salvo no Drive (falha).\n',
                     'O arquivo foi enviado no Telegram acima.\n',
                     'Para reenviar depois, responda "manda o pdf".']
        await query.edit_message_text('\n'.join(lines))


async def _gerar_e_salvar(context, fid, flow):
    """Generates parecer, DOCX+PDF, uploads to Drive, sends to Telegram."""
    from telegram import InlineKeyboardButton, InlineKeyboardMarkup, InputFile
    email_data = flow['email_data']
    empresa = email_data.get('empresa', 'Cliente')
    cnpj = email_data.get('cnpj', '')
    modelo = flow['modelo']

    try:
        parecer_text = generate_parecer_claude_v2(email_data, modelo)
        if not parecer_text:
            await context.bot.send_message(chat_id=CHAT_ID, text='❌ Erro ao gerar parecer (API Claude).')
            return
        flow['parecer_text'] = parecer_text

        ts = datetime.now().strftime('%Y%m%d_%H%M')
        empresa_clean = re.sub(r'[^\w\s\-]', '', empresa).strip()[:40]
        base_filename = f'PARECER_{empresa_clean}_{ts}'
        flow['base_filename'] = base_filename

        docx_buf = generate_parecer_docx(parecer_text, empresa, cnpj, modelo)
        docx_bytes_raw = docx_buf.read() if docx_buf else None

        # Upload ao Drive como Google Doc (sem cota de storage para service accounts)
        gdoc_url = None
        pdf_bytes = None
        if docx_bytes_raw:
            gdoc_url, pdf_bytes, drive_err = upload_to_drive_as_gdoc(
                docx_bytes_raw, base_filename, flow['folder_id']
            )
            if drive_err and not gdoc_url:
                await context.bot.send_message(
                    chat_id=CHAT_ID,
                    text=f'⚠️ Drive: {drive_err}'
                )

        # Fallback: gera PDF localmente via fpdf2 se Drive falhou
        if not pdf_bytes:
            try:
                pdf_bytes = generate_parecer_pdf_bytes(parecer_text, empresa, cnpj, modelo)
            except Exception as e:
                await context.bot.send_message(chat_id=CHAT_ID, text=f'⚠️ fpdf2: {e}')

        flow['docx_url'] = gdoc_url   # Link do Google Doc no Drive
        flow['pdf_url'] = None        # PDF exportado em memória, não salvo separado
        flow['pdf_bytes'] = pdf_bytes
        flow['state'] = 'await_auth'

        # Salva globalmente + Redis para reenvio posterior (inclui docx_bytes como fallback)
        persist_last_parecer({
            'base_filename': base_filename,
            'pdf_bytes': pdf_bytes,
            'docx_bytes': docx_bytes_raw,
            'pdf_url': pdf_url,
            'docx_url': docx_url,
            'empresa': empresa,
            'email_data': email_data,
        })

        # Envia PDF (ou DOCX como fallback) antes de mostrar os botões
        if pdf_bytes:
            await context.bot.send_document(
                chat_id=CHAT_ID,
                document=InputFile(_io_module.BytesIO(pdf_bytes), filename=base_filename + '.pdf'),
                caption=f'📑 {base_filename}.pdf — revise antes de autorizar o envio'
            )
        elif docx_bytes_raw:
            await context.bot.send_document(
                chat_id=CHAT_ID,
                document=InputFile(_io_module.BytesIO(docx_bytes_raw), filename=base_filename + '.docx'),
                caption=f'⚠️ PDF indisponível (geração falhou). Revise o DOCX antes de autorizar.'
            )

        lines = [
            f'✅ PARECER GERADO — {modelo}\n',
            f'🏢 {empresa}' + (f'  |  CNPJ: {cnpj}' if cnpj else ''),
            f'📁 Pasta: {flow["folder_name"]}\n',
        ]
        if gdoc_url:
            lines.append(f'📄 Google Doc (Drive): {gdoc_url}')
        lines.append('\nRevise o PDF acima e decida:')

        buttons = [
            [InlineKeyboardButton('✅ Autorizar envio por email', callback_data=f'as:{fid}')],
            [InlineKeyboardButton('💾 Salvar só no Drive', callback_data=f'ac:{fid}')],
        ]
        await context.bot.send_message(
            chat_id=CHAT_ID,
            text='\n'.join(lines),
            reply_markup=InlineKeyboardMarkup(buttons)
        )

        _parecer_pending.pop(flow.get('uid'), None)
        persist_parecer_pending()

    except Exception as e:
        await context.bot.send_message(chat_id=CHAT_ID, text=f'❌ Erro ao processar parecer: {e}')


async def _enviar_email_resposta(query, context, fid, flow):
    email_data = flow['email_data']
    pdf_bytes = flow.get('pdf_bytes')
    base_filename = flow.get('base_filename', 'PARECER')

    if not pdf_bytes:
        await query.edit_message_text('❌ PDF não disponível para envio.')
        return

    summary = (
        f'Prezado(a),\n\n'
        f'Segue em anexo o parecer técnico da Compliance-CE '
        f'em resposta à sua consulta: {email_data["subject"]}.\n\n'
        f'Qualquer dúvida, estamos à disposição.\n\n'
        f'Atenciosamente,\n'
        f'Marcos Lima — CRC/CE\n'
        f'Compliance-CE'
    )

    ok, err = reply_email_with_pdf(email_data, pdf_bytes, base_filename, summary)
    if ok:
        _parecer_flow.pop(fid, None)
        await query.edit_message_text(
            f'✅ Email enviado!\nPara: {email_data["sender"]}\nAnexo: {base_filename}.pdf'
        )
    else:
        await query.edit_message_text(f'❌ Falha ao enviar email:\n{(err or "")[:300]}')


@only_authorized
async def cmd_remetentes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lista os remetentes com mais emails não lidos, agrupados por quantidade."""
    text = update.message.text.lower()
    # Detecta se pediu filtro de tempo
    if 'hoje' in text or '24h' in text:
        hours, periodo = 24, 'hoje'
    elif 'semana' in text:
        hours, periodo = 168, 'semana'
    else:
        hours, periodo = None, 'caixa toda'

    await update.message.reply_text(f'🔍 Mapeando remetentes ({periodo})...')
    result, spam_count, err = fetch_by_sender(hours=hours, limit=300)
    if err:
        await update.message.reply_text(f'❌ Erro: {err}')
        return
    if not result:
        await update.message.reply_text('📭 Nenhum email não lido.')
        return

    total_emails = sum(c for _, _, c, _ in result)
    lines = [f'👥 REMETENTES — {len(result)} únicos, {total_emails} emails ({periodo})\n']
    for nome, addr, count, cat in result[:25]:
        icon = ICONS.get(cat, '📌')
        bar = '█' * min(count, 10) + ('…' if count > 10 else '')
        lines.append(f'{icon} {bar} {count}x  {nome}\n   {addr}')

    if len(result) > 25:
        lines.append(f'\n... e mais {len(result) - 25} remetentes')
    lines.append(f'\n🗑️ Spam/propaganda ignorado: {spam_count}')

    await update.message.reply_text('\n'.join(lines))


@only_authorized
async def cmd_pessoal_gmail(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Lê somente a conta pessoal cnp.marcoslima@gmail.com"""
    pessoal = next((a for a in ACCOUNTS if 'cnp' in a['user']), None)
    if not pessoal:
        await update.message.reply_text('❌ Conta pessoal não configurada.')
        return
    await update.message.reply_text(f'🔍 Lendo {pessoal["user"]}...')
    try:
        mail = get_imap(pessoal['user'], pessoal['password'], readonly=True)
        ids = search_unseen(mail, hours=None)
        total = len(ids)
        buckets = {k: [] for k in ICONS}
        spam_count = 0
        for eid, sender, subject in fetch_headers(mail, ids, limit=100):
            if is_spam(sender, subject):
                spam_count += 1
                continue
            cat = classify(sender, subject)
            name = sender.split('<')[0].strip().strip('"') or sender
            name = name[:35] + '...' if len(name) > 35 else name
            sub = subject[:55] + '...' if len(subject) > 55 else subject
            buckets[cat].append(f'• {sub}\n  ↳ {name}')
        mail.logout()
        msg = build_message(buckets, f'👤 Pessoal ({pessoal["user"]}) — {total} não lidos, {spam_count} propagandas')
        await update.message.reply_text(msg)
    except Exception as e:
        await update.message.reply_text(f'❌ Erro: {e}')


@only_authorized
async def cmd_varredura(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('🔍 Varrendo caixa de entrada...')
    try:
        total = 0
        encontrados = []
        for acc in ACCOUNTS:
            mail = get_imap(acc['user'], acc['password'], readonly=True)
            ids = search_unseen(mail, hours=None)
            total += len(ids)
            for eid, sender, subject in fetch_headers(mail, ids, limit=500):
                kw = match_alert(sender, subject)
                if kw:
                    name = sender.split('<')[0].strip().strip('"') or sender
                    name = name[:35] + '...' if len(name) > 35 else name
                    sub = subject[:55] + '...' if len(subject) > 55 else subject
                    encontrados.append((kw, sub, name, acc['label']))
            mail.logout()

        if not encontrados:
            await update.message.reply_text(
                f'✅ Varredura concluída — {total} emails analisados\n'
                f'Nenhum email com palavras-chave críticas encontrado.'
            )
            return

        lines = [f'🚨 VARREDURA — {len(encontrados)} alertas em {total} emails\n']
        for kw, sub, name, label in encontrados[:30]:
            lines.append(f'⚡ [{kw.upper()}] {label}\n• {sub}\n  ↳ {name}\n')
        msg = '\n'.join(lines)
        await update.message.reply_text(msg)
        if len(encontrados) > 30:
            await update.message.reply_text(f'... e mais {len(encontrados)-30} emails com alertas.')
    except Exception as e:
        await update.message.reply_text(f'❌ Erro na varredura: {e}')


@only_authorized
async def cmd_alertas(update: Update, context: ContextTypes.DEFAULT_TYPE):
    kws = '\n'.join(f'• {k}' for k in sorted(ALERT_KEYWORDS))
    await update.message.reply_text(f'🚨 Palavras-chave monitoradas ({len(ALERT_KEYWORDS)}):\n\n{kws}')


@only_authorized
async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        '🤖 Fale naturalmente ou use comandos:\n\n'
        '📊 "ler os 1000 emails" — resumo geral\n'
        '☀️ "o que chegou hoje?" — últimas 24h\n'
        '📅 "e essa semana?" — últimos 7 dias\n'
        '🔴 "tem algo urgente?" — urgentes\n'
        '💰 "quanto vendi?" — vendas Eduzz/Hotmart\n'
        '👨‍💼 "ver profissional" — jurídico/tributário\n'
        '💳 "ver financeiro" — XP, Pix, banco\n'
        '👦 "email do davi?" — pessoal/escola\n'
        '👥 "quem mais manda email?" — ranking por remetente\n'
        '🚨 "ver alertas" — palavras monitoradas\n'
        '✅ "zera a caixa" — marca tudo como lido\n\n'
        '🔔 Monitoramento automático ativo a cada 10 min!\n'
        'Alertas disparam para: clientes, SPED, PGFN, etc.\n\n'
        '📝 PARECERES:\n'
        'Quando chegar email pedindo orientação/análise tributária,\n'
        'o bot avisa automaticamente. Responda "parecer" para gerar\n'
        'o parecer preliminar com as 10 seções da Compliance.'
    )


async def _execute_tool(tool_name, tool_input, update, context):
    """Executa a ferramenta escolhida pelo Claude e retorna texto do resultado."""
    global _last_shown_email
    periodo_map = {'hoje': 24, 'semana': 168, 'tudo': None}

    if tool_name == 'ver_emails':
        periodo = tool_input.get('periodo', 'hoje')
        cat = tool_input.get('categoria')
        hours = periodo_map.get(periodo, 24)
        buckets, total, spam, err = fetch_emails(hours=hours, only_cat=cat, limit=100)
        if err:
            await update.message.reply_text(f'❌ Erro: {err}')
        else:
            title = f'📬 {periodo.capitalize()} — {total} total, {spam} propagandas'
            await update.message.reply_text(build_message(buckets, title))

    elif tool_name == 'contar_emails':
        total, counts, spam, err = count_by_category(limit=500)
        if err:
            await update.message.reply_text(f'❌ Erro: {err}')
        else:
            lines = [f'📊 {total} emails não lidos\n']
            for cat, n in counts.items():
                if n > 0:
                    lines.append(f'{ICONS[cat]} {cat}: {n}')
            lines.append(f'🗑️ Propaganda: {spam}')
            await update.message.reply_text('\n'.join(lines))

    elif tool_name == 'buscar_email':
        nome = tool_input['nome']
        await update.message.reply_text(f'🔍 Buscando email de {nome}...')
        data = search_email_by_name(nome)
        if not data:
            await update.message.reply_text(f'📭 Nenhum email encontrado de "{nome}".')
        else:
            persist_last_email(data)
            await update.message.reply_text(
                f'📧 De: {data["sender"]}\n'
                f'📌 Assunto: {data["subject"]}\n'
                f'📅 Data: {data["date"]}\n\n'
                f'{data["body"][:800]}'
            )

    elif tool_name == 'analisar_email':
        nome = tool_input['nome']
        await update.message.reply_text(f'🔍 Buscando email de {nome}...')
        data = search_email_by_name(nome)
        if not data:
            await update.message.reply_text(f'📭 Nenhum email encontrado de "{nome}".')
        else:
            persist_last_email(data)
            await update.message.reply_text('⏳ Analisando com IA...')
            analysis = analyze_with_claude(data)
            await update.message.reply_text(analysis)

    elif tool_name == 'gerar_parecer':
        await cmd_gerar_parecer(update, context)

    elif tool_name == 'listar_remetentes':
        periodo = tool_input.get('periodo', 'tudo')
        hours = periodo_map.get(periodo)
        await update.message.reply_text(f'🔍 Mapeando remetentes ({periodo})...')
        result, spam_count, err = fetch_by_sender(hours=hours, limit=300)
        if err:
            await update.message.reply_text(f'❌ {err}')
        else:
            total_emails = sum(c for _, _, c, _ in result)
            lines = [f'👥 {len(result)} remetentes, {total_emails} emails ({periodo})\n']
            for nome, addr, count, cat in result[:20]:
                icon = ICONS.get(cat, '📌')
                bar = '█' * min(count, 10) + ('…' if count > 10 else '')
                lines.append(f'{icon} {bar} {count}x  {nome}\n   {addr}')
            if len(result) > 20:
                lines.append(f'\n... e mais {len(result) - 20} remetentes')
            await update.message.reply_text('\n'.join(lines))

    elif tool_name == 'varredura_alertas':
        await cmd_varredura(update, context)

    elif tool_name == 'marcar_lidos':
        total = mark_as_read()
        await update.message.reply_text(
            f'✅ {total} email(s) marcados como lidos.' if total >= 0 else '❌ Erro ao marcar emails.'
        )

    elif tool_name == 'reenviar_pdf':
        from telegram import InputFile
        p = get_last_parecer()
        if not p:
            await update.message.reply_text('📭 Nenhum parecer encontrado. Gere um novo com "elabora o parecer".')
            return
        base_fn = p.get('base_filename', 'PARECER')
        sent_bytes = None
        if p.get('pdf_bytes'):
            await context.bot.send_document(
                chat_id=CHAT_ID,
                document=InputFile(_io_module.BytesIO(p['pdf_bytes']), filename=base_fn + '.pdf'),
                caption=f'📑 {base_fn}.pdf'
            )
            sent_bytes = p['pdf_bytes']
        elif p.get('docx_bytes'):
            # PDF não disponível — envia DOCX como fallback
            await context.bot.send_document(
                chat_id=CHAT_ID,
                document=InputFile(_io_module.BytesIO(p['docx_bytes']), filename=base_fn + '.docx'),
                caption=f'⚠️ PDF indisponível. Aqui está o DOCX do parecer: {base_fn}'
            )
        elif p.get('pdf_url'):
            await update.message.reply_text(f'📑 PDF no Drive (não tenho os bytes localmente):\n{p["pdf_url"]}')
        elif p.get('docx_url'):
            await update.message.reply_text(f'📄 DOCX no Drive (PDF indisponível):\n{p["docx_url"]}')
        else:
            await update.message.reply_text('📭 Não há arquivo disponível para este parecer. Gere um novo com "elabora o parecer".')
            return
        lines = []
        if p.get('pdf_url'):
            lines.append(f'📑 Drive PDF: {p["pdf_url"]}')
        if p.get('docx_url'):
            lines.append(f'📄 Drive DOCX: {p["docx_url"]}')
        if lines:
            await update.message.reply_text('\n'.join(lines))
        # Envia por email se solicitado e temos bytes de PDF
        if tool_input.get('enviar_email') and p.get('email_data') and sent_bytes and p.get('pdf_bytes'):
            summary = (
                f'Prezado(a),\n\nSegue o parecer técnico da Compliance-CE '
                f'referente a: {p["email_data"]["subject"]}.\n\n'
                f'Atenciosamente,\nMarcos Lima — CRC/CE\nCompliance-CE'
            )
            ok, err = reply_email_with_pdf(p['email_data'], p['pdf_bytes'], base_fn, summary)
            await update.message.reply_text(
                f'✅ Email enviado para {p["email_data"]["sender"]}' if ok
                else f'❌ Falha ao enviar email: {err}'
            )


@only_authorized
async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Intercepta digitação de nome de pasta durante fluxo de parecer
    for fid, flow in list(_parecer_flow.items()):
        if flow.get('state') == 'folder_new':
            from telegram import InlineKeyboardButton, InlineKeyboardMarkup
            folder_name = update.message.text.strip()
            await update.message.reply_text(f'⏳ Criando pasta "{folder_name}"...')
            folder_id, _ = create_drive_folder(folder_name, DRIVE_PARECER_ROOT_ID)
            if folder_id:
                flow['folder_id'] = folder_id
                flow['folder_name'] = folder_name
                flow['state'] = 'modelo_select'
                buttons = [
                    [InlineKeyboardButton('📋 MODELO GERAL', callback_data=f'mg:{fid}')],
                    [InlineKeyboardButton('🏢 MODELO EMPRESARIAL', callback_data=f'me:{fid}')],
                ]
                await update.message.reply_text(
                    f'✅ Pasta criada: {folder_name}\n\n📝 Selecione o modelo do parecer:',
                    reply_markup=InlineKeyboardMarkup(buttons)
                )
            else:
                await update.message.reply_text('❌ Erro ao criar pasta. Verifique o acesso ao Drive.')
            return

    # Comandos diretos que não precisam de IA
    text_lower = update.message.text.lower().strip()
    if text_lower in ('ajuda', 'help', '/ajuda', '/help'):
        await cmd_help(update, context)
        return

    # Sem Claude configurado — fallback para help
    if not claude_client:
        await update.message.reply_text('❌ ANTHROPIC_API_KEY não configurada.')
        return

    # ── Claude como cérebro ──────────────────────────────────────────────────
    user_text = update.message.text.strip()
    context_note = ''
    if _parecer_pending:
        n = len(_parecer_pending)
        context_note += f'\n\n[Contexto: {n} solicitação(ões) de parecer pendente(s) na fila.]'
    if _last_shown_email:
        context_note += (
            f'\n\n[Último email exibido: '
            f'De "{_last_shown_email["sender"]}" | '
            f'Assunto: "{_last_shown_email["subject"]}". '
            f'Se o usuário pedir parecer sem especificar outro email, use este como base.]'
        )
    if _last_parecer:
        context_note += (
            f'\n\n[Último parecer gerado: "{_last_parecer["base_filename"]}" '
            f'para {_last_parecer["empresa"]}. '
            f'Se o usuário quiser ver o PDF ou enviar por email, use a ferramenta reenviar_pdf.]'
        )

    try:
        response = claude_client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=800,
            system=AGENT_SYSTEM + context_note,
            tools=CLAUDE_TOOLS,
            messages=[{'role': 'user', 'content': user_text}]
        )

        if response.stop_reason == 'tool_use':
            for block in response.content:
                if block.type == 'tool_use':
                    await _execute_tool(block.name, block.input, update, context)
                    break
        else:
            reply = ''.join(b.text for b in response.content if hasattr(b, 'text')).strip()
            if reply:
                await update.message.reply_text(reply)

    except Exception as e:
        await update.message.reply_text(f'❌ Erro ao processar: {e}')


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    load_state_from_redis()

    app = Application.builder().token(TELEGRAM_TOKEN).build()

    # Monitoramento automático a cada 10 minutos
    app.job_queue.run_repeating(check_alerts, interval=600, first=30)

    from telegram.ext import CallbackQueryHandler
    app.add_handler(CallbackQueryHandler(handle_callback_query))

    app.add_handler(CommandHandler('start', cmd_todos))
    app.add_handler(CommandHandler('todos', cmd_todos))
    app.add_handler(CommandHandler('emails', cmd_emails))
    app.add_handler(CommandHandler('hoje', cmd_hoje))
    app.add_handler(CommandHandler('semana', cmd_semana))
    app.add_handler(CommandHandler('urgente', cmd_urgente))
    app.add_handler(CommandHandler('vendas', cmd_vendas))
    app.add_handler(CommandHandler('profissional', cmd_profissional))
    app.add_handler(CommandHandler('financeiro', cmd_financeiro))
    app.add_handler(CommandHandler('pessoal', cmd_pessoal))
    app.add_handler(CommandHandler('analisar', cmd_analisar_email))
    app.add_handler(CommandHandler('parecer', cmd_gerar_parecer))
    app.add_handler(CommandHandler('remetentes', cmd_remetentes))
    app.add_handler(CommandHandler('cnp', cmd_pessoal_gmail))
    app.add_handler(CommandHandler('varredura', cmd_varredura))
    app.add_handler(CommandHandler('alertas', cmd_alertas))
    app.add_handler(CommandHandler('marcar', cmd_marcar_lido))
    app.add_handler(CommandHandler('ajuda', cmd_help))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    print('Bot rodando com monitoramento a cada 10 minutos...')
    app.run_polling(drop_pending_updates=True)


if __name__ == '__main__':
    main()
